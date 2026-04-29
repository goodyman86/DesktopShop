from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    run = heading.runs[0] if heading.runs else heading.add_run(text)
    if level == 0:
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    elif level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x16, 0x21, 0x3e)
    elif level == 3:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x0f, 0x3c, 0x78)
    return heading

def add_table_from_md(doc, lines):
    rows = []
    for line in lines:
        if re.match(r'\s*\|[-| :]+\|\s*$', line):
            continue
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        rows.append(cells)
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = 'Table Grid'
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j >= col_count:
                break
            cell = table.cell(i, j)
            cell.text = cell_text
            para = cell.paragraphs[0]
            run = para.runs[0] if para.runs else para.add_run(cell_text)
            run.font.size = Pt(9)
            if i == 0:
                run.bold = True
                set_cell_bg(cell, '1a1a2e')
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            elif i % 2 == 0:
                set_cell_bg(cell, 'EEF2FF')
    doc.add_paragraph()

def add_code_block(doc, lines):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.3)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F4F4F4')
    pPr.append(shd)
    text = '\n'.join(lines)
    run = para.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    doc.add_paragraph()

def parse_inline(para, text):
    bold_pattern = re.compile(r'\*\*(.+?)\*\*')
    code_pattern = re.compile(r'`(.+?)`')
    combined = re.compile(r'(\*\*(.+?)\*\*|`(.+?)`)')
    last = 0
    for m in combined.finditer(text):
        if m.start() > last:
            para.add_run(text[last:m.start()])
        if m.group(0).startswith('**'):
            run = para.add_run(m.group(2))
            run.bold = True
        else:
            run = para.add_run(m.group(3))
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xc7, 0x25, 0x4e)
        last = m.end()
    if last < len(text):
        para.add_run(text[last:])

def convert(md_path, docx_path):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Default font
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)

    with open(md_path, encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    in_table = False
    table_lines = []
    in_code = False
    code_lines = []

    while i < len(lines):
        line = lines[i].rstrip('\n')

        # Code block
        if line.strip().startswith('```'):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                add_code_block(doc, code_lines)
                in_code = False
                code_lines = []
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # Table detection
        if '|' in line and line.strip().startswith('|'):
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            continue
        else:
            if in_table:
                add_table_from_md(doc, table_lines)
                in_table = False
                table_lines = []

        stripped = line.strip()

        # Headings
        if stripped.startswith('# ') and not stripped.startswith('## '):
            add_heading(doc, stripped[2:], level=0)
        elif stripped.startswith('## '):
            add_heading(doc, stripped[3:], level=1)
        elif stripped.startswith('### '):
            add_heading(doc, stripped[4:], level=2)
        elif stripped.startswith('#### '):
            add_heading(doc, stripped[5:], level=3)

        # Horizontal rule
        elif stripped.startswith('---'):
            para = doc.add_paragraph()
            pPr = para._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), '1a1a2e')
            pBdr.append(bottom)
            pPr.append(pBdr)

        # Bullet list
        elif stripped.startswith('- '):
            para = doc.add_paragraph(style='List Bullet')
            parse_inline(para, stripped[2:])
            para.paragraph_format.left_indent = Inches(0.25)
            for run in para.runs:
                run.font.size = Pt(10)

        # Numbered list
        elif re.match(r'^\d+\. ', stripped):
            para = doc.add_paragraph(style='List Number')
            parse_inline(para, re.sub(r'^\d+\. ', '', stripped))
            para.paragraph_format.left_indent = Inches(0.25)
            for run in para.runs:
                run.font.size = Pt(10)

        # Bold line (metadata at top)
        elif stripped.startswith('**') and stripped.endswith('**') and len(stripped) > 4:
            para = doc.add_paragraph()
            run = para.add_run(stripped.strip('*'))
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        # Empty line
        elif stripped == '':
            pass

        # Normal paragraph
        else:
            if stripped.startswith('> '):
                stripped = stripped[2:]
            para = doc.add_paragraph()
            parse_inline(para, stripped)
            for run in para.runs:
                run.font.size = Pt(10.5)

        i += 1

    if in_table:
        add_table_from_md(doc, table_lines)

    doc.save(docx_path)
    print(f"Đã xuất: {docx_path}")

convert('/home/user/DesktopShop/BAOCAO.md', '/home/user/DesktopShop/BAOCAO.docx')
